from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Dict

from monitoring_app.config import APP_NAME, EXPORTS_DIR, FONT_CANDIDATES, OWNER_NAME
from monitoring_app.storage import DatabaseManager
from monitoring_app.utils.text import compact_text, safe_slug


class ReportService:
    def __init__(self, repository: DatabaseManager) -> None:
        self.repository = repository

    def generate_report(self, format_name: str, category: str = "", minimum_risk: int = 0) -> Path:
        rows = self.repository.export_rows(category=category, minimum_risk=minimum_risk)
        if rows.empty:
            raise ValueError("لا توجد بيانات مطابقة لتوليد التقرير.")

        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        report_slug = safe_slug(f"{APP_NAME}-{format_name}-{timestamp}")
        path = EXPORTS_DIR / f"{report_slug}.{format_name.lower()}"

        if format_name.lower() == "csv":
            rows.to_csv(path, index=False, encoding="utf-8-sig")
        elif format_name.lower() == "json":
            records = json.loads(rows.fillna("").to_json(orient="records", force_ascii=False, date_format="iso"))
            payload = {
                "generated_at": datetime.utcnow().isoformat(),
                "category_filter": category,
                "minimum_risk": minimum_risk,
                "records": records,
            }
            path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        elif format_name.lower() == "docx":
            self._write_docx(path, rows, category, minimum_risk)
        elif format_name.lower() == "pdf":
            self._write_pdf(path, rows, category, minimum_risk)
        else:
            raise ValueError(f"صيغة غير مدعومة: {format_name}")

        self.repository.record_report(path.stem, format_name.upper(), str(path), {"category": category, "minimum_risk": minimum_risk})
        return path

    def _write_docx(self, path: Path, rows, category: str, minimum_risk: int) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt

        document = Document()
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        normal_style.font.size = Pt(11)

        title = document.add_heading(APP_NAME, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        subtitle = document.add_paragraph(
            f"الجهة المالكة: {OWNER_NAME} | الفئة: {category or 'الكل'} | الحد الأدنى للخطورة: {minimum_risk}"
        )
        subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        case_rows = rows.drop_duplicates(subset=["case_id"])
        for _, row in case_rows.iterrows():
            heading = document.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            heading.add_run(f"القضية #{int(row['case_id'])}: {row['case_title']}").bold = True

            details = document.add_paragraph()
            details.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            details.add_run(f"التصنيف: {row['primary_category']} | الخطورة: {int(row['case_risk_score'])}/100 | الأدلة: {int(row['evidence_count'])}")

            summary = document.add_paragraph(compact_text(str(row["case_summary"]), 700))
            summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.save(path)

    def _write_pdf(self, path: Path, rows, category: str, minimum_risk: int) -> None:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        font_name = self._ensure_font()
        document = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            rightMargin=1.2 * cm,
            leftMargin=1.2 * cm,
            topMargin=1.5 * cm,
            bottomMargin=1.5 * cm,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ArabicTitle",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=18,
            leading=24,
            alignment=2,
            textColor=colors.HexColor("#111111"),
        )
        body_style = ParagraphStyle(
            "ArabicBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10.5,
            leading=16,
            alignment=2,
            textColor=colors.HexColor("#111111"),
        )
        story = [
            Paragraph(self._shape(APP_NAME), title_style),
            Spacer(1, 0.25 * cm),
            Paragraph(self._shape(f"الجهة المالكة: {OWNER_NAME}"), body_style),
            Paragraph(self._shape(f"الفئة: {category or 'الكل'} | الحد الأدنى للخطورة: {minimum_risk}"), body_style),
            Spacer(1, 0.35 * cm),
        ]

        case_rows = rows.drop_duplicates(subset=["case_id"])
        summary_table = Table(
            [
                [self._shape("عدد القضايا"), self._shape("عدد السجلات"), self._shape("أعلى خطورة")],
                [
                    self._shape(str(case_rows.shape[0])),
                    self._shape(str(rows.shape[0])),
                    self._shape(str(int(case_rows["case_risk_score"].max()))),
                ],
            ],
            hAlign="RIGHT",
        )
        summary_table.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D4AF37")),
                    ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B4912E")),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("PADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.extend([summary_table, Spacer(1, 0.35 * cm)])

        for _, row in case_rows.iterrows():
            story.append(
                Paragraph(
                    self._shape(f"القضية #{int(row['case_id'])}: {row['case_title']}"),
                    body_style,
                )
            )
            story.append(
                Paragraph(
                    self._shape(
                        f"التصنيف: {row['primary_category']} | الخطورة: {int(row['case_risk_score'])}/100 | الأدلة: {int(row['evidence_count'])}"
                    ),
                    body_style,
                )
            )
            story.append(Paragraph(self._shape(compact_text(str(row["case_summary"]), 600)), body_style))
            story.append(Spacer(1, 0.25 * cm))

        document.build(story)

    def _ensure_font(self) -> str:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        for candidate in FONT_CANDIDATES:
            if candidate.exists():
                font_name = candidate.stem
                if font_name not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(font_name, str(candidate)))
                return font_name
        return "Helvetica"

    def _shape(self, text: str) -> str:
        try:
            import arabic_reshaper
            from bidi.algorithm import get_display

            return get_display(arabic_reshaper.reshape(text or ""))
        except Exception:
            return text or ""
